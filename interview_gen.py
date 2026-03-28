"""
interview_gen.py  –  Generate targeted interview questionnaires.

Produces role-specific, skill-gap-aware questions based on:
  1. JD requirements vs candidate skills (gap analysis)
  2. Agent panel evaluations (each agent's focus questions)
  3. Verification findings (areas needing clarification)
  4. Behavioural/STAR-format questions tailored to experience level
"""
from __future__ import annotations

from dataclasses import dataclass, field
from typing import Optional

from models import CandidateProfile, CandidateScore, JDCriteria

# Try importing agent types  –  optional dependency
try:
    from agents import ConsensusResult
except ImportError:
    ConsensusResult = None


@dataclass
class InterviewQuestion:
    category: str           # "Technical" / "Behavioural" / "Security" / "Design" / "Situational"
    question: str
    rationale: str          # why this question matters for this candidate
    difficulty: str         # "Basic" / "Intermediate" / "Advanced"
    target_skill: str       # which skill/gap this probes
    suggested_followup: str


@dataclass
class InterviewQuestionnaire:
    candidate_name: str
    role_title: str
    generated_at: str
    total_questions: int
    sections: dict[str, list[InterviewQuestion]]   # category → questions
    skill_gaps_addressed: list[str]
    summary: str


# ---------------------------------------------------------------------------
# Question Templates
# ---------------------------------------------------------------------------

_TECHNICAL_TEMPLATES = {
    "aws": [
        ("How would you design a multi-account AWS landing zone for this organisation?", "Advanced"),
        ("Explain the difference between ECS and EKS. When would you choose one over the other?", "Intermediate"),
        ("Walk me through your approach to AWS cost optimisation.", "Intermediate"),
    ],
    "azure": [
        ("How would you implement a hub-spoke network architecture in Azure?", "Advanced"),
        ("Explain Azure Policy vs RBAC. How do they complement each other?", "Intermediate"),
        ("Describe your experience with Azure DevOps pipelines and deployment strategies.", "Intermediate"),
    ],
    "kubernetes": [
        ("Explain how you would troubleshoot a pod stuck in CrashLoopBackOff.", "Intermediate"),
        ("How would you implement zero-downtime deployments in Kubernetes?", "Advanced"),
        ("Describe your approach to Kubernetes RBAC and network policies.", "Advanced"),
    ],
    "terraform": [
        ("How do you manage Terraform state in a team environment?", "Intermediate"),
        ("Explain Terraform module composition patterns you've used.", "Advanced"),
        ("How do you handle secrets in Terraform configurations?", "Intermediate"),
    ],
    "docker": [
        ("How do you optimise Docker image sizes for production?", "Intermediate"),
        ("Explain multi-stage Docker builds and when you'd use them.", "Basic"),
        ("How do you handle security scanning in your container pipeline?", "Intermediate"),
    ],
    "ci/cd": [
        ("Design a CI/CD pipeline for a microservices application. What stages would you include?", "Advanced"),
        ("How do you implement rollback strategies in your deployment pipeline?", "Intermediate"),
        ("Explain your approach to testing in CI/CD (unit, integration, e2e).", "Intermediate"),
    ],
    "python": [
        ("Explain Python's GIL and its implications for concurrent programming.", "Advanced"),
        ("How do you structure a Python project for maintainability?", "Intermediate"),
        ("Describe your experience with Python async/await patterns.", "Intermediate"),
    ],
    "monitoring": [
        ("How would you set up observability for a distributed system?", "Advanced"),
        ("Explain the difference between metrics, logs, and traces.", "Basic"),
        ("How do you define SLOs and error budgets?", "Intermediate"),
    ],
    "security": [
        ("How do you implement secrets management in a cloud environment?", "Intermediate"),
        ("Explain your approach to zero-trust security architecture.", "Advanced"),
        ("How do you integrate security scanning into CI/CD?", "Intermediate"),
    ],
    "microservices": [
        ("How do you handle distributed transactions across microservices?", "Advanced"),
        ("Explain service mesh patterns and when you'd implement one.", "Advanced"),
        ("How do you approach API versioning in a microservices architecture?", "Intermediate"),
    ],
}

_BEHAVIOURAL_TEMPLATES = [
    ("Tell me about a time you had to make a critical technical decision under pressure. What was the outcome?",
     "decision-making", "Intermediate"),
    ("Describe a situation where you disagreed with a team member on a technical approach. How did you resolve it?",
     "collaboration", "Basic"),
    ("Tell me about the most complex system you helped design or build. What were the trade-offs?",
     "system design", "Advanced"),
    ("Describe a production incident you handled. What was your approach and what did you learn?",
     "incident response", "Intermediate"),
    ("How do you stay current with new technologies? Give an example of something you learned recently and applied.",
     "learning", "Basic"),
    ("Tell me about a time you had to onboard to a new codebase quickly. What strategy did you use?",
     "adaptability", "Basic"),
    ("Describe a time you mentored a junior team member. What was the challenge?",
     "leadership", "Intermediate"),
    ("Tell me about a project where requirements changed significantly mid-stream. How did you adapt?",
     "agility", "Intermediate"),
]

_SITUATIONAL_TEMPLATES = {
    "Junior": [
        ("If you encountered a bug in production that you caused, what steps would you take?",
         "accountability"),
        ("How would you approach learning a new technology required for your role?",
         "learning"),
    ],
    "Mid": [
        ("A stakeholder asks for an urgent feature that would require technical debt. How do you handle it?",
         "prioritisation"),
        ("You notice a colleague's code has a security vulnerability in a PR. How do you address it?",
         "communication"),
    ],
    "Senior": [
        ("You're asked to reduce infrastructure costs by 30%. Walk me through your approach.",
         "strategy"),
        ("Two teams need conflicting architectural decisions. How do you facilitate resolution?",
         "leadership"),
    ],
    "Lead": [
        ("How would you build a platform engineering team from scratch?",
         "team building"),
        ("Describe how you'd create a technical roadmap for the next 12 months.",
         "vision"),
    ],
}


# ---------------------------------------------------------------------------
# Generator
# ---------------------------------------------------------------------------

def generate_questionnaire(
    candidate: CandidateProfile,
    jd: JDCriteria,
    score: CandidateScore,
    consensus: object = None,  # ConsensusResult
) -> InterviewQuestionnaire:
    """Generate a complete interview questionnaire for a candidate."""
    from datetime import datetime

    sections: dict[str, list[InterviewQuestion]] = {}
    skill_gaps_addressed: list[str] = []

    # --- 1. Technical Questions (from skill gaps) ---
    technical_qs = []

    # Missing required skills → targeted questions
    for skill in score.missing_required_skills:
        sk_lower = skill.lower()
        skill_gaps_addressed.append(skill)

        # Check if we have templates for this skill
        matched_template = None
        for key, templates in _TECHNICAL_TEMPLATES.items():
            if key in sk_lower or sk_lower in key:
                matched_template = templates
                break

        if matched_template:
            for q_text, diff in matched_template[:2]:
                technical_qs.append(InterviewQuestion(
                    category="Technical",
                    question=q_text,
                    rationale=f"Candidate is missing required skill: {skill}",
                    difficulty=diff,
                    target_skill=skill,
                    suggested_followup=f"Can you provide a specific example of using {skill} in production?",
                ))
        else:
            technical_qs.append(InterviewQuestion(
                category="Technical",
                question=f"Describe your experience with {skill}. How have you used it in your projects?",
                rationale=f"Required skill '{skill}' not found in resume",
                difficulty="Intermediate",
                target_skill=skill,
                suggested_followup=f"How would you approach learning {skill} if you haven't used it before?",
            ))

    # Role-specific technical depth questions
    resume_lower = candidate.raw_text.lower()
    for key, templates in _TECHNICAL_TEMPLATES.items():
        if key in resume_lower:
            # Candidate claims this skill — probe depth
            q_text, diff = templates[-1]  # use hardest question
            technical_qs.append(InterviewQuestion(
                category="Technical",
                question=q_text,
                rationale=f"Candidate claims {key} experience — validating depth",
                difficulty=diff,
                target_skill=key,
                suggested_followup=f"What's a common pitfall with {key} that you've encountered?",
            ))

    sections["Technical"] = technical_qs[:8]

    # --- 2. Behavioural Questions ---
    behavioural_qs = []
    for q_text, target, diff in _BEHAVIOURAL_TEMPLATES:
        # Select based on experience level
        if candidate.total_experience_years < 3 and diff == "Advanced":
            continue
        if candidate.total_experience_years > 8 and diff == "Basic":
            continue

        behavioural_qs.append(InterviewQuestion(
            category="Behavioural",
            question=q_text,
            rationale=f"Assessing {target} for {jd.role_level} level role",
            difficulty=diff,
            target_skill=target,
            suggested_followup="What would you do differently if you faced the same situation today?",
        ))

    sections["Behavioural"] = behavioural_qs[:4]

    # --- 3. Situational Questions (level-appropriate) ---
    level = jd.role_level if jd.role_level in _SITUATIONAL_TEMPLATES else "Mid"
    situational_qs = []
    for q_text, target in _SITUATIONAL_TEMPLATES.get(level, _SITUATIONAL_TEMPLATES["Mid"]):
        situational_qs.append(InterviewQuestion(
            category="Situational",
            question=q_text,
            rationale=f"Role-level ({level}) situational assessment: {target}",
            difficulty="Intermediate",
            target_skill=target,
            suggested_followup="How would the outcome change in a different organisational context?",
        ))
    sections["Situational"] = situational_qs

    # --- 4. Agent-driven Questions (from consensus) ---
    if consensus and hasattr(consensus, "evaluations"):
        agent_qs = []
        for ev in consensus.evaluations:
            for q in ev.key_questions[:1]:  # top question per agent
                agent_qs.append(InterviewQuestion(
                    category=f"Agent: {ev.agent_name}",
                    question=q,
                    rationale=f"{ev.agent_name} perspective — {ev.persona} focus",
                    difficulty="Intermediate",
                    target_skill=ev.persona.lower(),
                    suggested_followup="",
                ))
        if agent_qs:
            sections["Agent Panel"] = agent_qs

    # --- 5. Verification follow-up questions ---
    verification_qs = []
    ver = score.verification
    if ver:
        # Unverified companies
        unverified = [c.name for c in ver.companies if not c.found]
        if unverified:
            verification_qs.append(InterviewQuestion(
                category="Verification",
                question=f"We couldn't verify {'these companies' if len(unverified) > 1 else 'this company'}: "
                         f"{', '.join(unverified[:3])}. Can you provide more details about your role there?",
                rationale="Company verification failed — need candidate clarification",
                difficulty="Basic",
                target_skill="experience verification",
                suggested_followup="Can you provide a reference from this company?",
            ))

        # Timeline gaps
        if ver.identity and ver.identity.timeline and ver.identity.timeline.has_gaps:
            for gap in ver.identity.timeline.gap_details[:2]:
                verification_qs.append(InterviewQuestion(
                    category="Verification",
                    question=f"There appears to be a gap in your timeline: {gap}. "
                             f"Can you tell us about this period?",
                    rationale="Employment timeline gap detected",
                    difficulty="Basic",
                    target_skill="timeline verification",
                    suggested_followup="",
                ))

    if verification_qs:
        sections["Verification"] = verification_qs

    # --- Build questionnaire ---
    total = sum(len(qs) for qs in sections.values())

    summary_parts = [
        f"Interview questionnaire for **{candidate.name}** targeting the **{jd.title}** role ({jd.role_level}).",
        f"Contains {total} questions across {len(sections)} sections.",
    ]
    if skill_gaps_addressed:
        summary_parts.append(f"Addresses {len(skill_gaps_addressed)} skill gaps: {', '.join(skill_gaps_addressed[:5])}.")

    return InterviewQuestionnaire(
        candidate_name=candidate.name,
        role_title=jd.title,
        generated_at=datetime.now().isoformat(),
        total_questions=total,
        sections=sections,
        skill_gaps_addressed=skill_gaps_addressed,
        summary=" ".join(summary_parts),
    )


# ---------------------------------------------------------------------------
# DOCX Export
# ---------------------------------------------------------------------------

def export_questionnaire_docx(q: InterviewQuestionnaire) -> bytes:
    """Generate a professional DOCX document from an InterviewQuestionnaire.

    Returns the DOCX file contents as bytes (ready for download).
    """
    from docx import Document
    from docx.shared import Pt, Inches, Cm, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_TABLE_ALIGNMENT
    import io

    doc = Document()

    # Page margins
    for section in doc.sections:
        section.top_margin = Cm(2)
        section.bottom_margin = Cm(2)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)

    # Default font
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Calibri'
    font.size = Pt(11)
    font.color.rgb = RGBColor(0x33, 0x33, 0x33)

    # Title
    title = doc.add_heading('Interview Questionnaire', level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in title.runs:
        run.font.color.rgb = RGBColor(0x1a, 0x1a, 0x2e)

    # Subtitle
    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run(f'{q.role_title}')
    run.font.size = Pt(16)
    run.font.color.rgb = RGBColor(0x3b, 0x82, 0xf6)
    run.bold = True

    # Candidate info table
    doc.add_paragraph()
    info_table = doc.add_table(rows=4, cols=2)
    info_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    info_data = [
        ('Candidate', q.candidate_name),
        ('Role', q.role_title),
        ('Generated', q.generated_at[:10] if len(q.generated_at) >= 10 else q.generated_at),
        ('Total Questions', str(q.total_questions)),
    ]
    for i, (label, value) in enumerate(info_data):
        cell_label = info_table.cell(i, 0)
        cell_value = info_table.cell(i, 1)
        cell_label.text = label
        cell_value.text = value
        for paragraph in cell_label.paragraphs:
            for run in paragraph.runs:
                run.bold = True
                run.font.size = Pt(10)
        for paragraph in cell_value.paragraphs:
            for run in paragraph.runs:
                run.font.size = Pt(10)

    # Skill gaps
    if q.skill_gaps_addressed:
        doc.add_paragraph()
        p = doc.add_paragraph()
        run = p.add_run('Skill Gaps to Probe: ')
        run.bold = True
        run.font.size = Pt(11)
        run.font.color.rgb = RGBColor(0xef, 0x44, 0x44)
        run = p.add_run(', '.join(q.skill_gaps_addressed))
        run.font.size = Pt(11)

    doc.add_paragraph()

    # Sections
    _DIFFICULTY_COLORS = {
        'Basic': RGBColor(0x10, 0xb9, 0x81),
        'Intermediate': RGBColor(0xf5, 0x9e, 0x0b),
        'Advanced': RGBColor(0xef, 0x44, 0x44),
    }

    for section_name, questions in q.sections.items():
        # Section header
        heading = doc.add_heading(f'{section_name} ({len(questions)} questions)', level=1)
        for run in heading.runs:
            run.font.color.rgb = RGBColor(0x0f, 0x34, 0x60)

        for i, qq in enumerate(questions, 1):
            # Question number + difficulty badge
            p_q = doc.add_paragraph()
            run_num = p_q.add_run(f'Q{i}. ')
            run_num.bold = True
            run_num.font.size = Pt(11)

            run_diff = p_q.add_run(f'[{qq.difficulty}] ')
            run_diff.bold = True
            run_diff.font.size = Pt(9)
            run_diff.font.color.rgb = _DIFFICULTY_COLORS.get(qq.difficulty, RGBColor(0x94, 0xa3, 0xb8))

            run_question = p_q.add_run(qq.question)
            run_question.font.size = Pt(11)

            # Target skill
            if qq.target_skill:
                p_target = doc.add_paragraph()
                p_target.paragraph_format.left_indent = Cm(1)
                run = p_target.add_run(f'Target: {qq.target_skill}')
                run.italic = True
                run.font.size = Pt(9)
                run.font.color.rgb = RGBColor(0x64, 0x74, 0x8b)

            # Rationale
            if qq.rationale:
                p_rat = doc.add_paragraph()
                p_rat.paragraph_format.left_indent = Cm(1)
                run = p_rat.add_run(f'Rationale: {qq.rationale}')
                run.italic = True
                run.font.size = Pt(9)
                run.font.color.rgb = RGBColor(0x94, 0xa3, 0xb8)

            # Follow-up
            if qq.suggested_followup:
                p_fu = doc.add_paragraph()
                p_fu.paragraph_format.left_indent = Cm(1)
                run = p_fu.add_run(f'Follow-up: {qq.suggested_followup}')
                run.italic = True
                run.font.size = Pt(9)
                run.font.color.rgb = RGBColor(0x94, 0xa3, 0xb8)

            # Notes area (blank line for interviewer)
            p_notes = doc.add_paragraph()
            p_notes.paragraph_format.left_indent = Cm(1)
            run = p_notes.add_run('Notes: ____________________________________________________________')
            run.font.size = Pt(9)
            run.font.color.rgb = RGBColor(0xbd, 0xbd, 0xbd)

    # Footer
    doc.add_paragraph()
    footer = doc.add_paragraph()
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = footer.add_run(f'Generated by Resume Screener Pro • {q.generated_at[:10]}')
    run.font.size = Pt(8)
    run.font.color.rgb = RGBColor(0x94, 0xa3, 0xb8)
    run.italic = True

    # Save to bytes
    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.getvalue()
